import pandas as pd
import numpy as np
import warnings
import os
import docx
import re
import nltk

from nltk.tokenize import sent_tokenize
nltk.download('punkt')

warnings.filterwarnings('ignore')


def makemydir(current_wd, new_directory):
    """
    Creates a new directory if it doesn't exist in the current working directory.
    current_wd: current working directory
    new_directory: name of the new directory to be created
    """
    os.chdir(current_wd)
    try:
        os.makedirs(new_directory)
    except OSError:
        pass
    # let exception propagate if we just can't
    # cd into the specified directory
    os.chdir(new_directory)
    
# --------------------------------------------- DataFrame Machine Translator ---------------------------------------------
def machine_translator_split(target_language, source_language, current_wd, df, column_name, Translator):
    """
    Splits the column to be translated into batches, according to the translator we could have different limits in characters, e.g., google translator limit is 5,000 in one request.
    Stores the batches in CSVs.
    Input
    current_wd: working directory (location for storing the batches)
    df: Data Frame that will be translated
    column_name: Column to be translated

    Output:
    Location of the csv files
    """     
    max_allowed = 5000     # translator character limit (depends on the chosen translator)
    
    # substract by languagea
    df['original_index'] = df.index

    count_name = 'Count_' + column_name
    cumsm_name = 'CumSm_' + column_name
    batch_name = 'Batch_' + column_name
    
    try:
        df[count_name] = df[column_name].str.len()
    except KeyError:
        raise KeyError(f"No column with the name '{column_name}'")

    df[cumsm_name] = df[count_name].cumsum()
    
    sub_df = df[['original_index', column_name, count_name, cumsm_name]]
    sub_df[batch_name] = np.floor(sub_df[cumsm_name]/max_allowed) + 1

    max_batch = np.max(sub_df[batch_name])
    print('\nMaximum number of batches:', int(max_batch))
    
    folder_name = column_name + '_' + target_language
    makemydir(current_wd, folder_name)

    for i in range(1, int(max_batch)+1):
        # Percentage done
        print('Translating batch', i, '...', round( (i/max_batch)*100, 2), '%')

        # Filter information to translate
        source_to_translate = sub_df[sub_df[batch_name] == i][['original_index', column_name]]
        text_to_translate = source_to_translate[column_name]

        # Translate
        translated_df = Translator(source=source_language, target=target_language).translate_batch(text_to_translate.tolist())
        
        # Gather translation
        new_column_name = column_name + '_' + target_language
        stored_translated = pd.DataFrame({new_column_name: translated_df})

        # Information gathering
        finished_batch_translation = pd.concat([source_to_translate.reset_index(drop=True), stored_translated], axis=1)

        # Batch file name
        store_name = column_name + '_' + 'batch_' + str(i) + '.csv'

        # Store
        finished_batch_translation.to_csv(store_name, encoding='utf-16', index=False, sep = '|')
        
    stored_location = os.getcwd()
    os.chdir(current_wd)

    print('\nTotal characters translated:', int(np.max(sub_df[cumsm_name])) )
    print('Translations stored in:', stored_location)
    return stored_location



def merge_csvs(target_language, stored_location, df, column_name):
    """
    Reads all the csv files from "machine_translation" function with the original dataframe
    Input
    stored_location: working directory (location for storing the batches)
    df: Original dataframe that was translated
    column_name: Column that was translated

    Output:
    Original dataframe with a new column (translation)
    """     
    # Translations output path
    files = os.listdir(os.path.join(stored_location))
    
    # Merge all translations
    combined_csv = pd.concat([pd.read_csv( os.path.join(stored_location, f) , encoding='utf-16', sep='|') for f in files ])
    # Clean the index
    combined_csv.index = combined_csv['original_index']
    combined_csv.drop('original_index', axis=1, inplace=True)
    combined_csv.sort_index(inplace=True)
    trans_complete = combined_csv.drop([column_name], axis=1)
    # Clean df
    df.drop(['original_index', 'Count_'+ column_name, 'CumSm_' + column_name], axis=1, inplace=True)

    all_df = df.merge(trans_complete,left_index=True,right_index=True, how='left')
    new_column_name = column_name + '_' + target_language
    all_df = all_df.loc[all_df[new_column_name].notnull()]

    print('Number of files read:', len(files))
    return all_df

    
def machine_translator_df(data_set, column_name, target_language, source_language, Translator, current_wd):
    # -------------------- Translate datset by batches (stored in multiple .csv) ---------------------
    stored_location = machine_translator_split(target_language, source_language, current_wd, data_set, column_name, Translator)
    print('Dataset was divided by batches located in:', stored_location)
    # ----------------------------- Merge the translated .csv files -----------------------------------
    df_merged_translation = merge_csvs(target_language, stored_location, data_set, column_name)

    return df_merged_translation


def store_translation(data_set, output_path, output_name):
    """
    Stores the new dataframe
    Input
    data_set: Complete translated dataframe
    output_path: Path of the new file 
    output_name: Name of the new fill

    Output:
    Original dataframe with a new column (translation)
    """   

    # csv
    output_name_without_extension = output_name.replace(".csv", "")

    output_final = os.path.join(output_path, output_name_without_extension + '.csv')
    data_set.to_csv(output_final, encoding='utf-16', index=False)
    print('Final stored in:', output_final, '\n')


# --------------------------------------------- Document Machine Translator ---------------------------------------------
def read_word_document(root, data_set_name):
    """
    This function reads a Word document (.docx) located at a given directory and returns the text contents as a single string.

    Args:
        root (str): The path to the directory where the document is located.
        data_set_name (str): The filename of the document to be read.

    Returns:
        str: A string containing the text contents of the Word document.

    """
    # Construct the full file path by joining the specified root directory and document filename
    file_path = os.path.join(root, data_set_name)
    # Open the Word document using the python-docx library
    document = docx.Document(file_path)
    # Initialize an empty list to hold each paragraph of text in the document
    text = []
    # Iterate through each paragraph in the document and append its text to the list
    for paragraph in document.paragraphs:
        text.append(paragraph.text)
        text.append('\n')

    return "\n".join(text) # Join the list of paragraph texts into a single string and return it

def save_text_to_docx(root, text, output_filename):
    """
    This function saves a given text string as a new Word document (.docx) with the specified filename.

    Args:
        root (str): The path to the directory where the new document should be saved.
        text (str): The text to be saved in the new document.
        output_filename (str): The desired filename for the new document.

    Returns:
        None.

    """
    # Create a new Word document and add a new paragraph containing the input text
    document = docx.Document()
    document.add_paragraph(text)
    # Construct the full output path by joining the specified root directory and output filename
    full_output_path = os.path.join(root, output_filename)
    # Save the new Word document to the specified output path
    document.save(full_output_path)


def replace_consecutive_chars(text, pattern='-'):
    """
    This function replaces consecutive occurrences of a character or sequence of characters in a string with newline characters.

    Args:
        text (str): The input string to be processed.
        pattern (str): The character or sequence of characters to be replaced (default is '-' for hyphens).

    Returns:
        str: The processed string with newline characters replacing consecutive occurrences of the specified pattern.

    """
    # Construct a regular expression pattern that matches 2 or more consecutive occurrences of the specified pattern
    pattern = f'{pattern}{{1,}}' # match 2 or more consecutive hyphens
    
    return re.sub(pattern, '\n', text) # Replace all occurrences of the pattern with a single newline character



def machine_translator_doc(text, target_language, source_language, Translator, current_wd):
    """
    This function takes in text data stored in a Word document and translates it into the target language using the specified machine translation service provider.

    Args:
        text (str): The text data from the Word document.
        target_language (str): The target language for translation.      
        source_language
        Translator (object): The machine translation service provider to use.
        current_wd (str): The root directory where the Word document is stored.

    Returns:
        str: The translated text.

    Raises:
        KeyError: If the specified column name is not present in the DataFrame.
    """
        
    # Replace newline characters with hyphens and tokenize the text into sentences
    s_text = text.replace('\n', ' -')
    sentences = sent_tokenize(s_text)

    # Convert it to df in order to use the function "machine_translator_split"
    df = pd.DataFrame(sentences, columns=['Sentence'])
    # Call the machine_translator_split function to perform the translation
    stored_location= machine_translator_split(target_language, source_language, current_wd, df, 'Sentence', Translator)

    # Merge the translated CSV files into a complete translated DataFrame
    df_merged_translation = merge_csvs(target_language, stored_location, df, 'Sentence')

    # Join the translated sentences with spaces and replace consecutive hyphens with newlines
    translated_text = " ".join(df_merged_translation['Sentence_' + target_language])
    #replace_consecutive
    result=replace_consecutive_chars(translated_text ,'-')

    return result
